from flask import Flask, render_template, request, send_file
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm ,Inches, RGBColor
from datetime import datetime
from copy import deepcopy
from flask import after_this_request
import os
import uuid
app = Flask(__name__)

# ==========================================================
# ROAD TABLE
# ==========================================================
def insert_road_table(doc, marker_text, road_data):

    for paragraph in doc.paragraphs:

        if marker_text in paragraph.text:

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            parent.remove(paragraph._element)

            num_types = len(road_data)
            max_rows = max(v["LHS"] + v["RHS"] for v in road_data.values())

            table = doc.add_table(rows=max_rows + 1, cols=num_types)
            table.style = "Table Grid"

            # header row
            for col, road_type in enumerate(road_data.keys()):

                cell = table.cell(0, col)
                cell.text = road_type

                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # fill rows
            for col, (road_type, sides) in enumerate(road_data.items()):

                row_index = 1

                abbr = "".join(word[0] for word in road_type.split()).upper()
                display_name = road_type.rstrip("s")

                for i in range(1, sides["LHS"] + 1):

                    table.cell(row_index, col).text = f"{abbr}L {i} – {display_name} LHS {i}"
                    row_index += 1

                for i in range(1, sides["RHS"] + 1):

                    table.cell(row_index, col).text = f"{abbr}R {i} – {display_name} RHS {i}"
                    row_index += 1

            parent.insert(index, table._element)

            break


# ==========================================================
# EXECUTIVE SUMMARY TABLE COPY
# ==========================================================
def copy_executive_summary_table(source_doc, target_doc, marker_text):

    for paragraph in target_doc.paragraphs:

        if marker_text in paragraph.text:

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            parent.remove(paragraph._element)

            found_reference = False

            for block in source_doc.element.body:

                if block.tag.endswith("p"):

                    text = "".join(node.text for node in block.iter() if node.text)

                    if "summary of Gap study report" in text:
                        found_reference = True
                        continue

                if found_reference and block.tag.endswith("tbl"):

                    parent.insert(index, deepcopy(block))
                    break

            break 

# fix size of tables
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

def fix_table_width(table):
    tbl = table._element
    tblPr = tbl.tblPr

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(8500))  # ~15 cm

    tblPr.append(tblW)


def fix_table_layout(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()

        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), "0")
        trPr.append(cantSplit)


def reduce_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

# ==========================================================
# GENERIC SECTION EXTRACTION
# ==========================================================


from docx.shared import Inches
import os

def insert_section(target_doc, source_doc, marker_text, start_heading, stop_heading):
    """
    Copy section content from source_doc to target_doc between start_heading and stop_heading.
    Supports:
    - paragraphs
    - tables
    - images (graphs) inside paragraphs
    """

    capture = False

    for paragraph in target_doc.paragraphs:
        if marker_text in paragraph.text:

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            parent.remove(paragraph._element)

            for block in source_doc.element.body:

                # Detect paragraph text
                if block.tag.endswith("p"):
                    text = "".join(node.text for node in block.iter() if node.text).strip()

                    if start_heading in text:
                        capture = True
                        continue

                    if capture and stop_heading in text:
                        break

                if capture:

                    # If table
                    if block.tag.endswith("tbl"):
                        parent.insert(index, deepcopy(block))
                        index += 1
                        continue

                    # If paragraph
                    if block.tag.endswith("p"):

                        # Look for image reference in this paragraph
                        rId = None
                        for node in block.iter():
                            if node.tag.endswith("blip"):
                                rId = node.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")

                        if rId:
                            # Extract correct image
                            image_part = source_doc.part.related_parts[rId]
                            image_bytes = image_part.blob

                            img_path = f"temp_image_{rId}.png"

                            with open(img_path, "wb") as f:
                                f.write(image_bytes)

                            p = target_doc.add_paragraph()
                            run = p.add_run()
                            run.add_picture(img_path)

                            parent.insert(index, p._element)
                            index += 1
                            os.remove(img_path)

                        else:
                            parent.insert(index, deepcopy(block))
                            index += 1

            break


# ==========================================================
# SECTION TABLE-ONLY EXTRACTION
# ==========================================================
def insert_section_tables_only(target_doc, source_doc, marker_text, start_heading, stop_heading):
    """
    Copy ONLY tables from source_doc (between start_heading and stop_heading)
    into target_doc at the marker_text placeholder.
    """

    capture = False

    for paragraph in target_doc.paragraphs:
        if marker_text in paragraph.text:

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            parent.remove(paragraph._element)

            for block in source_doc.element.body:

                # Detect paragraph text for heading matching
                if block.tag.endswith("p"):
                    text = "".join(node.text for node in block.iter() if node.text).strip()

                    if start_heading in text:
                        capture = True
                        continue

                    if capture and stop_heading in text:
                        break

                if capture:
                    # Only copy tables, skip everything else
                    if block.tag.endswith("tbl"):
                        parent.insert(index, deepcopy(block))
                        index += 1

            break
# COPY FIRST TABLE AFTER HEADING 4.2
def copy_first_table_after_heading(source_doc, target_doc, marker_text, heading_text):

    for paragraph in target_doc.paragraphs:
        if marker_text in paragraph.text:

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            found_heading = False

            for block in source_doc.element.body:

                # detect heading
                if block.tag.endswith("p"):
                    text = "".join(node.text for node in block.iter() if node.text).strip()

                    if heading_text in text:
                        found_heading = True
                        continue

                # get first table after heading
                if found_heading and block.tag.endswith("tbl"):

                    parent.insert(index, deepcopy(block))
                    parent.remove(paragraph._element)
                    break

            break


# COPY FIRST IMAGE AFTER HEADING CHAINAGE WISE ANALYSED DOCUMENT FOR PLACEHOLDER RESULT_GAP_STUDY_TWO_GRAPH_ONE
def copy_first_image_after_main_heading(source_doc, target_doc, marker_text, heading_text):

    for paragraph in target_doc.paragraphs:
        if marker_text in paragraph.text:

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            found_heading = False

            for block in source_doc.element.body:

                # Detect heading
                if block.tag.endswith("p"):
                    text = "".join(node.text for node in block.iter() if node.text).strip()

                    if heading_text.lower() in text.lower():
                        found_heading = True
                        continue

                # After heading → search for first image
                if found_heading and block.tag.endswith("p"):

                    for node in block.iter():
                        if node.tag.endswith("blip"):

                            rId = node.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )

                            image_part = source_doc.part.related_parts[rId]
                            image_bytes = image_part.blob

                            img_path = f"temp_chainage_graph.png"

                            with open(img_path, "wb") as f:
                                f.write(image_bytes)

                            # insert image
                            p = target_doc.add_paragraph()
                            run = p.add_run()
                            run.add_picture(img_path)

                            parent.insert(index, p._element)
                            parent.remove(paragraph._element)

                            return  # stop after first image


#  FOR PLACEHOLDER RESULT_GAP_STUDY_TWO_GRAPH_ONE, COPIES FIRST GRAPH AFTER HEADING "Chainage Wise Gap Analysis"
def copy_graph_after_table(source_doc, target_doc, marker_text, heading_text):

    for paragraph in target_doc.paragraphs:
        if marker_text in paragraph.text:

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            found_heading = False
            table_passed = False

            for block in source_doc.element.body:

                # detect heading
                if block.tag.endswith("p"):
                    text = "".join(node.text for node in block.iter() if node.text).strip()

                    if heading_text in text:
                        found_heading = True
                        continue

                if found_heading:

                    # detect table first
                    if block.tag.endswith("tbl"):
                        table_passed = True
                        continue

                    # after table → find image
                    if table_passed and block.tag.endswith("p"):

                        for node in block.iter():
                            if node.tag.endswith("blip"):

                                rId = node.get(
                                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                )

                                image_part = source_doc.part.related_parts[rId]
                                image_bytes = image_part.blob

                                img_path = f"temp_4_2_graph.png"

                                with open(img_path, "wb") as f:
                                    f.write(image_bytes)

                                # insert image
                                p = target_doc.add_paragraph()
                                run = p.add_run()
                                run.add_picture(img_path)

                                parent.insert(index, p._element)
                                parent.remove(paragraph._element)

                                return

# RSA_SUMMARY TABLE COPY
from copy import deepcopy

def insert_rsa_summary_table(source_doc, target_doc, marker_text):

    tables = source_doc.tables

    if not tables:
        print("No tables found")
        return

    # Step 1: Start from last table
    merged_tables = []
    merged_tables.append(tables[-1])

    # Step 2: Go backwards and collect connected tables
    for i in range(len(tables) - 2, -1, -1):

        current = tables[i]
        next_table = tables[i + 1]

        # Compare column count → same table continuation
        if len(current.columns) == len(next_table.columns):
            merged_tables.insert(0, current)
        else:
            break

    print(f"Tables merged: {len(merged_tables)}")

    # Step 3: Merge all tables
    base = merged_tables[0]
    merged_element = deepcopy(base._element)

    for table in merged_tables[1:]:
        for row in table.rows:
            merged_element.append(deepcopy(row._element))

    # Step 4: Insert into target doc
    for paragraph in target_doc.paragraphs:
        if marker_text in paragraph.text:

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            parent.insert(index, merged_element)
            parent.remove(paragraph._element)

            print("Full table inserted successfully")
            return

# fix_table width to fit page (for executive summary table and RSA summary table)
def fix_table_width(table):

    # Auto fit OFF (important)
    table.autofit = False

    # total page usable width (approx)
    total_width = Cm(16)   # adjust if needed

    num_cols = len(table.columns)
    col_width = total_width / num_cols

    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width
# ADD ANNEXURE_A
from copy import deepcopy

def insert_full_document(source_doc, target_doc, marker_text):

    for paragraph in target_doc.paragraphs:
        if marker_text in paragraph.text:

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            # Insert ALL content
            for element in source_doc.element.body:
                parent.insert(index, deepcopy(element))
                index += 1

            parent.remove(paragraph._element)

            print("Annexure content inserted successfully")
            return
# EXTRACT ANNEXURE D FROM ANALYSED DOCUMENT
def extract_till_end(source_doc, start_heading):

    content = []
    capture = False

    for block in source_doc.element.body:

        # detect heading
        if block.tag.endswith("p"):
            text = "".join(node.text for node in block.iter() if node.text).strip()

            if start_heading in text:
                capture = True
                continue

        if capture:
            content.append(deepcopy(block))

    return content


def insert_section_blocks(doc, marker_text, content_blocks, source_doc):

    for paragraph in doc.paragraphs:
        if marker_text in paragraph.text:

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            for block in content_blocks:

                # ======================
                # PARAGRAPH
                # ======================
                if block.tag.endswith("p"):

                    new_para = doc.add_paragraph()

                    for node in block.iter():

                        # TEXT
                        if node.tag.endswith("t") and node.text:
                            new_para.add_run(node.text)

                        # IMAGE
                        if node.tag.endswith("blip"):
                            rId = node.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')

                            if rId in source_doc.part.rels:

                                image_part = source_doc.part.rels[rId].target_part
                                image_bytes = image_part.blob

                                img_path = f"temp_{uuid.uuid4().hex}.png"

                                with open(img_path, "wb") as f:
                                    f.write(image_bytes)

                                run = new_para.add_run()
                                run.add_picture(img_path, width=Cm(14.4))

                                os.remove(img_path)

                    parent.insert(index, new_para._element)
                    index += 1

                # ======================
                # TABLE
                # ======================
                elif block.tag.endswith("tbl"):
                    new_table = deepcopy(block)
                    parent.insert(index, new_table)
                    index += 1

                    # convert to python-docx table object
                    table = doc.tables[-1]

                    fix_table_width(table)
                    fix_table_layout(table)
                    reduce_table_font(table)

            parent.remove(paragraph._element)
            break

# APPLY ZONE COLOR


def apply_zone_color(doc, zone):

    zone_colors = {
        "Zone A": RGBColor(255, 0, 0),     # Red
        "Zone B": RGBColor(0, 0, 255),     # Blue
        "Zone C": RGBColor(0, 128, 0),     # Green
        "Zone D": RGBColor(255, 255, 0),   # Yellow
        "Zone E": RGBColor(128, 0, 128)    # Purple
    }

    if zone not in zone_colors:
        return

    color = zone_colors[zone]

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if zone in run.text:
                run.font.color.rgb = color
                run.bold = True   # optional
# ==========================================================
# MAIN ROUTE
# ==========================================================
@app.route("/", methods=["GET", "POST"])
def form():

    if request.method == "POST":

        template_path = "template.docx"
        temp_path = "temp_rendered.docx"

        tpl = DocxTemplate(template_path)

        # DATE
        raw_date = request.form.get("starting_survey_date")
        formatted_date = ""

        if raw_date:
            formatted_date = datetime.strptime(
                raw_date, "%Y-%m-%d"
            ).strftime("%d %B %Y")

        # MAP IMAGE
        image = None
        image_file = request.files.get("map_image")

        if image_file and image_file.filename:

            image_path = "temp_map.png"
            image_file.save(image_path)

            image = InlineImage(tpl, image_path, width=Cm(12))
        # Create road_data
        selected_types = request.form.getlist("road_type")
        road_data = {}

        for road_type in selected_types:
            lhs = int(request.form.get(f"{road_type}_LHS", 0))
            rhs = int(request.form.get(f"{road_type}_RHS", 0))

            road_data[road_type] = {
                "LHS": lhs,
                "RHS": rhs
            }


        # Create road_summary
        road_summary = []

        for road_type, sides in road_data.items():
            total = sides["LHS"] + sides["RHS"]

            if total > 0:
                road_summary.append({
                    "name": road_type,
                    "total": total
                })

        # TEMPLATE CONTEXT
        context = {
            "project_name": request.form.get("project_name"),
            "upc_code": request.form.get("upc_code"),
            "state": request.form.get("state"),
            "ro": request.form.get("ro"),
            "piu": request.form.get("piu"),
            "length": request.form.get("length"),
            "flexibleorrigid": request.form.get("flexibleorrigid"),
            "lanes": request.form.get("lanes"),
            "om_dlp": request.form.get("om_dlp"),
            "starting_survey_date": formatted_date,
            "zone": request.form.get("zone"),
            "map_image": image,
            "road_summary": road_summary
        }
 

        tpl.render(context)
        tpl.save(temp_path)

        doc = Document(temp_path)

        # =========================================
        # ANALYZED DOCUMENT
        # =========================================
        uploaded_file = request.files.get("analysed_doc")

        if uploaded_file and uploaded_file.filename:

            analysed_path = f"uploaded_{datetime.now().strftime('%H%M%S')}.docx"
            uploaded_file.save(analysed_path)

            source_doc = Document(analysed_path)
             # ANNEXURE D
            # ==========================
            anx_d_content = extract_till_end(
                source_doc,
                "Chainage Wise Gap Analysis"
            )

            insert_section_blocks(
            doc,
            "### ANX_D ###",
            anx_d_content,
            source_doc   # IMPORTANT
          )




            # EXECUTIVE SUMMARY TABLE
            copy_executive_summary_table(
                source_doc,
                doc,
                "### INSERT_EXECUTIVE_SUMMARY_TABLE ###"
            )

            # SECTION 3.1
            insert_section(
                doc,
                source_doc,
                "### INSERT_INVENTORY_SECTION_ONE ###",
                "3.1",
                "3.2"
            )

            # SECTION 3.2
            insert_section(
                doc,
                source_doc,
                "### INSERT_INVENTORY_SECTION_TWO ###",
                "3.2",
                "3.3"
            )

            # SECTION 4.1 (full content)
            insert_section(
                doc,
                source_doc,
                "### RESULT_GAP_STUDY_ONE ###",
                "4.1",
                "4.2"
            )

            # SECTION 4.2 (tables only)
            insert_section_tables_only(
                doc,
                source_doc,
                "### RESULT_GAP_STUDY_TWO ###",
                "4.2",
                "4.3"
            )
            # SECTION 4.2 (first table after heading)
            copy_first_table_after_heading(
                source_doc,
                doc,
                "### RESULT_GAP_STUDY_TWO_TABLE ###",
                "4.2"
            )
            copy_first_image_after_main_heading(
                source_doc,
                doc,
                "### RESULT_GAP_STUDY_TWO_GRAPH_ONE ###",
                "Chainage Wise Gap Analysis"
            )
            copy_graph_after_table(
                source_doc,
                doc,
                "### RESULT_GAP_STUDY_TWO_GRAPH_TWO ###",
                "4.2"
            )
        # ==========================
        # ANNEXURE A
        # ==========================
        anx_a_file = request.files.get("anx_a_doc")

        if anx_a_file and anx_a_file.filename:
            path = "anx_a.docx"
            anx_a_file.save(path)
            anx_a_doc = Document(path)

            insert_full_document(anx_a_doc, doc, "### ANX_A ###")


        # ==========================
        # ANNEXURE B
        # ==========================
        anx_b_file = request.files.get("anx_b_doc")

        if anx_b_file and anx_b_file.filename:
            path = "anx_b.docx"
            anx_b_file.save(path)
            anx_b_doc = Document(path)

            insert_full_document(anx_b_doc,doc, "### ANX_B ###")


        # ==========================
        # ANNEXURE C
        # ==========================
        anx_c_file = request.files.get("anx_c_doc")

        if anx_c_file and anx_c_file.filename:
            path = "anx_c.docx"
            anx_c_file.save(path)
            anx_c_doc = Document(path)
            # INSERT FULL ANNEXURE C CONTENT
            insert_full_document(
                anx_c_doc,
                doc,
                "### ANX_C ###") 
            # INSERT RSA_SUMMARY TABLE (last table + connected tables)
            insert_rsa_summary_table(
                anx_c_doc,
                doc,
                "### RSA_SUMMARY ###"
                 ) 
            

        # =========================================
        # ROAD TABLE
        # =========================================
        selected_types = request.form.getlist("road_type")
        road_data = {}

        for road_type in selected_types:

            lhs = int(request.form.get(f"{road_type}_LHS", 0))
            rhs = int(request.form.get(f"{road_type}_RHS", 0))

            road_data[road_type] = {
                "LHS": lhs,
                "RHS": rhs
            }
        # -------------------------
        # ROAD SUMMARY
        # -------------------------
        road_summary = []

        for road_type, sides in road_data.items():
            total = sides["LHS"] + sides["RHS"]

            if total > 0:
                road_summary.append({
                    "name": road_type,
                    "total": total
                })
        if road_data:
            insert_road_table(doc, "### INSERT_ROAD_TABLE ###", road_data)
        # zone colour function
        apply_zone_color(doc, request.form.get("zone"))
        # SAVE FINAL REPORT
        project_name = request.form.get("project_name", "Report")

        safe_name = "".join(c for c in project_name if c.isalnum() or c in (" ", "_")).strip()

        final_path = f"Final_Report_{safe_name}.docx"

        doc.save(final_path)

        # ==========================
        # CLEANUP AFTER DOWNLOAD
        # ==========================
        @after_this_request
        def cleanup(response):
            try:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

                if os.path.exists("temp_map.png"):
                    os.remove("temp_map.png")

                for file in os.listdir():
                    if file.startswith("uploaded_") and file.endswith(".docx"):
                        os.remove(file)

                for file in os.listdir():
                    if file.startswith("temp_") and file.endswith(".png"):
                        os.remove(file)

            except Exception as e:
                print("Cleanup error:", e)

            return response

        # ✅ RETURN MUST BE HERE
        return send_file(final_path, as_attachment=True, download_name=final_path)

    return render_template("form.html")

if __name__ == "__main__":
    app.run(debug=True)