"""
Docx manipulation utilities — ported from newone/app.py (v2.1)
"""

import logging
import os
import uuid
from copy import deepcopy

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

log = logging.getLogger(__name__)

ZONE_COLORS = {
    "Zone A": RGBColor(255, 0, 0),
    "Zone B": RGBColor(0, 0, 255),
    "Zone C": RGBColor(0, 128, 0),
    "Zone D": RGBColor(255, 255, 0),
    "Zone E": RGBColor(128, 0, 128),
}

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


# ==========================================================
# XML ITERATION HELPERS
# ==========================================================
def _block_text(block) -> str:
    return "".join(n.text for n in block.iter() if n.text).strip()


def _iter_body(source_doc):
    for block in source_doc.element.body:
        tag = block.tag.split("}")[-1]
        yield tag, block


def _find_marker(target_doc, marker_text):
    for para in target_doc.paragraphs:
        if marker_text in para.text:
            parent = para._element.getparent()
            idx = list(parent).index(para._element)
            return parent, idx, para._element
    log.warning("Marker not found: %s", marker_text)
    return None, None, None


def _remove_marker(parent, marker_el):
    """
    Remove a marker paragraph while preserving any embedded sectPr.
    """
    pPr = marker_el.find(f"{{{W}}}pPr")
    sect_pr = pPr.find(f"{{{W}}}sectPr") if pPr is not None else None

    if sect_pr is not None:
        prev = marker_el.getprevious()
        while prev is not None and prev.tag != f"{{{W}}}p":
            prev = prev.getprevious()

        if prev is not None:
            prev_pPr = prev.find(f"{{{W}}}pPr")
            if prev_pPr is None:
                prev_pPr = OxmlElement("w:pPr")
                prev.insert(0, prev_pPr)
            old_sect = prev_pPr.find(f"{{{W}}}sectPr")
            if old_sect is not None:
                prev_pPr.remove(old_sect)
            prev_pPr.append(deepcopy(sect_pr))
            log.info("Preserved sectPr from marker to previous paragraph.")

    parent.remove(marker_el)


# ==========================================================
# TABLE SCALING / IMAGE TRANSFER
# ==========================================================
def _scale_table_in_block(block, scale: float = 1.1):
    for el in block.iter():
        local = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if local in ("tblW", "tcW", "gridCol"):
            w = el.get(f"{{{W}}}w")
            if w and w.lstrip("-").isdigit() and int(w) > 0:
                el.set(f"{{{W}}}w", str(int(int(w) * scale)))


def _scale_images_in_block(block, scale: float = 1.1):
    for el in block.iter():
        local = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if local == "extent":
            cx = el.get("cx")
            cy = el.get("cy")
            if cx and cx.lstrip("-").isdigit():
                el.set("cx", str(int(int(cx) * scale)))
            if cy and cy.lstrip("-").isdigit():
                el.set("cy", str(int(int(cy) * scale)))


def _transfer_image_rels(source_doc, dest_doc, block):
    from docx.parts.image import ImagePart
    block = deepcopy(block)
    REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"

    for blip in block.iter():
        if not blip.tag.endswith("blip"):
            continue
        rId = blip.get(f"{{{R}}}embed")
        if not rId or rId not in source_doc.part.rels:
            continue
        rel = source_doc.part.rels[rId]
        if "image" not in rel.reltype:
            continue
        try:
            src_img_part = rel.target_part
            new_img_part = ImagePart.load(
                src_img_part.partname,
                src_img_part.content_type,
                src_img_part.blob,
                dest_doc.part.package,
            )
            new_rId = dest_doc.part.relate_to(new_img_part, REL_TYPE)
            blip.set(f"{{{R}}}embed", new_rId)
        except Exception as e:
            log.error("Image transfer failed rId=%s: %s", rId, e)
    return block


# ==========================================================
# ROAD TABLE & EXECUTIVES
# ==========================================================
def insert_road_table(doc, marker_text: str, road_data: dict):
    parent, index, marker_el = _find_marker(doc, marker_text)
    if parent is None:
        return

    _remove_marker(parent, marker_el)
    max_rows = max((v["LHS"] + v["RHS"] for v in road_data.values()), default=0)
    if max_rows == 0:
        return

    num_types = len(road_data)
    table = doc.add_table(rows=max_rows + 1, cols=num_types)
    table.style = "Table Grid"

    for col, road_type in enumerate(road_data.keys()):
        cell = table.cell(0, col)
        cell.text = road_type
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for col, (road_type, sides) in enumerate(road_data.items()):
        abbr = "".join(w[0] for w in road_type.split()).upper()
        display = road_type.rstrip("s")
        row_index = 1
        for i in range(1, sides["LHS"] + 1):
            table.cell(row_index, col).text = f"{abbr}L {i} – {display} LHS {i}"
            row_index += 1
        for i in range(1, sides["RHS"] + 1):
            table.cell(row_index, col).text = f"{abbr}R {i} – {display} RHS {i}"
            row_index += 1

    parent.insert(index, table._element)
    reduce_table_font(table, 9)
    log.info("Road table inserted (%d types, %d rows)", num_types, max_rows)


def copy_executive_summary_table(source_doc, target_doc, marker_text: str):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return
    _remove_marker(parent, marker_el)
    found_ref = False
    for tag, block in _iter_body(source_doc):
        if tag == "p" and "summary of gap study report" in _block_text(block).lower():
            found_ref = True
            continue
        if found_ref and tag == "tbl":
            parent.insert(index, deepcopy(block))
            log.info("Executive summary table copied.")
            return
    log.warning("Executive summary table not found in source document.")


def extract_executive_summary_values(source_doc) -> dict:
    ROW_INDEX   = -1      
    ROW_KEYWORD = "total" 
    COL_NHAI    = 1
    COL_AI      = 2
    COL_COMBO   = 3
    COL_DIFF    = 4
    COL_RSA     = 5

    defaults = {
        "ai_survey_count": 0, "nhai_record_count": 0, "ai_nhai_diffrence": 0,
        "ai_and_rsa": 0, "rsa": 0, "excess_shortfall": "N/A",
    }

    target_table = None
    found_ref = False

    for tag, block in _iter_body(source_doc):
        if tag == "p" and "summary of gap study report" in _block_text(block).lower():
            found_ref = True
            continue
        if found_ref and tag == "tbl":
            from docx.table import Table as DocxTable
            target_table = DocxTable(block, source_doc)
            break

    if target_table is None:
        return defaults

    row = None
    if ROW_INDEX >= 0:
        if len(target_table.rows) > ROW_INDEX:
            row = target_table.rows[ROW_INDEX]
    else:
        for r in target_table.rows:
            if ROW_KEYWORD.lower() in r.cells[0].text.lower():
                row = r
                break

    if row is None:
        return defaults

    def _cell_int(cell_idx: int) -> int:
        try:
            raw = row.cells[cell_idx].text.strip()
            clean = raw.replace(",", "").replace("(", "-").replace(")", "")
            return int(clean)
        except (IndexError, ValueError):
            return 0

    nhai   = _cell_int(COL_NHAI)
    ai     = _cell_int(COL_AI)
    combo  = _cell_int(COL_COMBO)
    diff   = _cell_int(COL_DIFF)
    rsa    = _cell_int(COL_RSA)

    if ai > nhai:
        excess_shortfall = "Excess"
    elif ai < nhai:
        excess_shortfall = "Shortfall"
    else:
        excess_shortfall = "Nil"

    result = {
        "nhai_record_count":  nhai, "ai_survey_count":    ai, "ai_and_rsa":         combo,
        "ai_nhai_diffrence":  abs(diff), "rsa":                rsa, "excess_shortfall":   excess_shortfall,
    }
    return result


def reduce_table_font(table, size_pt: int = 9):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(size_pt)


# ==========================================================
# GENERIC SECTION EXTRACTION
# ==========================================================
def insert_section(target_doc, source_doc, marker_text: str, start_heading: str, stop_heading: str, image_scale: float = 1.15, table_scale: float = 1.1):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return
    _remove_marker(parent, marker_el)
    capture = False
    for tag, block in _iter_body(source_doc):
        if tag == "p":
            text = _block_text(block)
            if start_heading in text:
                capture = True
                continue
            if capture and stop_heading and stop_heading in text:
                break
        if not capture:
            continue
        if tag == "tbl":
            new_block = _transfer_image_rels(source_doc, target_doc, block)
            _scale_table_in_block(new_block, scale=table_scale)
            _strip_sect_pr(new_block)
            parent.insert(index, new_block)
            index += 1
        elif tag == "p":
            new_block = _transfer_image_rels(source_doc, target_doc, block)
            _scale_images_in_block(new_block, scale=image_scale)
            _strip_sect_pr(new_block)
            parent.insert(index, new_block)
            index += 1


def insert_section_tables_only(target_doc, source_doc, marker_text: str, start_heading: str, stop_heading: str, table_scale: float = 1.1):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return
    _remove_marker(parent, marker_el)
    capture = False
    for tag, block in _iter_body(source_doc):
        if tag == "p":
            text = _block_text(block)
            if start_heading in text:
                capture = True
                continue
            if capture and stop_heading and stop_heading in text:
                break
        if capture and tag == "tbl":
            new_block = _transfer_image_rels(source_doc, target_doc, block)
            _scale_table_in_block(new_block, scale=table_scale)
            _strip_sect_pr(new_block)
            parent.insert(index, new_block)
            index += 1


def copy_first_table_after_heading(source_doc, target_doc, marker_text: str, heading_text: str, table_scale: float = 1.1):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return
    found_heading = False
    for tag, block in _iter_body(source_doc):
        if tag == "p" and heading_text in _block_text(block):
            found_heading = True
            continue
        if found_heading and tag == "tbl":
            new_block = _transfer_image_rels(source_doc, target_doc, block)
            _scale_table_in_block(new_block, scale=table_scale)
            _strip_sect_pr(new_block)
            parent.insert(index, new_block)
            _remove_marker(parent, marker_el)
            return


def _extract_blip_rId(block):
    for node in block.iter():
        if node.tag.endswith("blip"):
            return node.get(f"{{{R}}}embed")
    return None


def copy_first_image_after_main_heading(source_doc, target_doc, marker_text: str, heading_text: str, image_scale: float = 1.15):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return
    found = False
    for tag, block in _iter_body(source_doc):
        if tag == "p" and heading_text.lower() in _block_text(block).lower():
            found = True
            continue
        if found and tag == "p":
            rId = _extract_blip_rId(block)
            if rId:
                new_block = _transfer_image_rels(source_doc, target_doc, block)
                _scale_images_in_block(new_block, scale=image_scale)
                _strip_sect_pr(new_block)
                parent.insert(index, new_block)
                _remove_marker(parent, marker_el)
                return


def copy_graph_after_table(source_doc, target_doc, marker_text: str, heading_text: str, image_scale: float = 1.15):
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return
    found_heading = False
    table_passed = False
    for tag, block in _iter_body(source_doc):
        if tag == "p" and heading_text in _block_text(block):
            found_heading = True
            continue
        if found_heading:
            if tag == "tbl":
                table_passed = True
                continue
            if table_passed and tag == "p":
                rId = _extract_blip_rId(block)
                if rId:
                    new_block = _transfer_image_rels(source_doc, target_doc, block)
                    _scale_images_in_block(new_block, scale=image_scale)
                    for sect_pr in new_block.findall(f".//{{{W}}}sectPr"):
                        sect_pr.getparent().remove(sect_pr)
                    parent.insert(index, new_block)
                    _remove_marker(parent, marker_el)
                    return

# ==========================================================
# HELPER FUNCTIONS FOR PAGE BREAKS & FORMATTING
# ==========================================================
def _strip_sect_pr(block):
    """Remove embedded sectPr and inject a manual page break to maintain pagination."""
    for sect_pr in block.findall(f".//{{{W}}}sectPr"):
        pPr = sect_pr.getparent()
        p = pPr.getparent() if pPr is not None else None
        pPr.remove(sect_pr)
        if p is not None and p.tag == f"{{{W}}}p":
            run = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br.set(f"{{{W}}}type", "page")
            run.append(br)
            p.append(run)

def _preserve_xml_spaces(block):
    """Forces lxml to preserve leading/trailing spaces in text runs (fixes PDF squishing)."""
    XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
    for t_el in block.iter(f"{{{W}}}t"):
        text = t_el.text or ""
        if text and (text != text.strip() or not text.strip()):
            t_el.set(XML_SPACE, "preserve")

def _center_tables_in_block(block):
    """Forces tables to center alignment and removes rogue left-indentation."""
    tables = [block] if block.tag.endswith("tbl") else block.findall(f".//{{{W}}}tbl")
    for tbl in tables:
        tbl_pr = tbl.find(f"{{{W}}}tblPr")
        if tbl_pr is None:
            tbl_pr = etree.SubElement(tbl, f"{{{W}}}tblPr")
            
        jc = tbl_pr.find(f"{{{W}}}jc")
        if jc is None:
            jc = etree.SubElement(tbl_pr, f"{{{W}}}jc")
        jc.set(f"{{{W}}}val", "center")
        
        tbl_ind = tbl_pr.find(f"{{{W}}}tblInd")
        if tbl_ind is None:
            tbl_ind = etree.SubElement(tbl_pr, f"{{{W}}}tblInd")
        tbl_ind.set(f"{{{W}}}w", "0")
        tbl_ind.set(f"{{{W}}}type", "dxa")

# ==========================================================
# RSA SUMMARY TABLE (merged + normalized)
# ==========================================================
def _fix_run_spacing(tbl_el):
    XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
    for t_el in tbl_el.iter(f"{{{W}}}t"):
        text = t_el.text or ""
        if not text.strip() or text != text.strip() or text == " ":
            t_el.set(XML_SPACE, "preserve")

def _normalize_table_width(tbl_el, target_doc):
    body = target_doc.element.body
    sect_pr = body.find(f".//{{{W}}}sectPr")
    page_w = 11906
    margin_left = 1440
    margin_right = 1440
    if sect_pr is not None:
        pg_sz = sect_pr.find(f"{{{W}}}pgSz")
        pg_mar = sect_pr.find(f"{{{W}}}pgMar")
        if pg_sz is not None: page_w = int(pg_sz.get(f"{{{W}}}w", page_w))
        if pg_mar is not None:
            margin_left = int(pg_mar.get(f"{{{W}}}left", margin_left))
            margin_right = int(pg_mar.get(f"{{{W}}}right", margin_right))
    content_width = page_w - margin_left - margin_right

    tbl_pr = tbl_el.find(f"{{{W}}}tblPr")
    if tbl_pr is None: tbl_pr = etree.SubElement(tbl_el, f"{{{W}}}tblPr")

    tbl_w_el = tbl_pr.find(f"{{{W}}}tblW")
    if tbl_w_el is None: tbl_w_el = etree.SubElement(tbl_pr, f"{{{W}}}tblW")
    tbl_w_el.set(f"{{{W}}}w", str(content_width))
    tbl_w_el.set(f"{{{W}}}type", "dxa")

    tbl_ind = tbl_pr.find(f"{{{W}}}tblInd")
    if tbl_ind is None: tbl_ind = etree.SubElement(tbl_pr, f"{{{W}}}tblInd")
    tbl_ind.set(f"{{{W}}}w", "0")
    tbl_ind.set(f"{{{W}}}type", "dxa")

    scale = 1.0
    tbl_grid = tbl_el.find(f"{{{W}}}tblGrid")
    if tbl_grid is not None:
        grid_cols = tbl_grid.findall(f"{{{W}}}gridCol")
        if grid_cols:
            original_total = sum(int(c.get(f"{{{W}}}w", 0)) for c in grid_cols)
            if original_total > 0 and original_total != content_width:
                scale = content_width / original_total
                for col in grid_cols:
                    orig = int(col.get(f"{{{W}}}w", 0))
                    col.set(f"{{{W}}}w", str(int(orig * scale)))

    for tc in tbl_el.iter(f"{{{W}}}tc"):
        tc_pr = tc.find(f"{{{W}}}tcPr")
        if tc_pr is None: continue
        tc_w = tc_pr.find(f"{{{W}}}tcW")
        if tc_w is not None:
            w_type = tc_w.get(f"{{{W}}}type", "dxa")
            if w_type == "dxa":
                orig = int(tc_w.get(f"{{{W}}}w", 0))
                if orig > 0: tc_w.set(f"{{{W}}}w", str(int(orig * scale)))

def _fix_header_row_text(tbl_el):
    rows = tbl_el.findall(f"{{{W}}}tr")
    if not rows: return
    header_row = rows[0]
    for tc in header_row.findall(f"{{{W}}}tc"):
        for p in tc.findall(f"{{{W}}}p"):
            tc.remove(p)
        new_p = etree.SubElement(tc, f"{{{W}}}p")
        pPr = etree.SubElement(new_p, f"{{{W}}}pPr")
        jc = etree.SubElement(pPr, f"{{{W}}}jc")
        jc.set(f"{{{W}}}val", "center")
        rPr_p = etree.SubElement(pPr, f"{{{W}}}rPr")
        etree.SubElement(rPr_p, f"{{{W}}}b")
        sz = etree.SubElement(rPr_p, f"{{{W}}}sz")
        sz.set(f"{{{W}}}val", "24")
        fonts = etree.SubElement(rPr_p, f"{{{W}}}rFonts")
        fonts.set(f"{{{W}}}ascii", "Arial")
        lines = ["ROAD SAFETY AUDITOR RECOMMENDATION FOR ROAD SIGNAGES", "DURING", "OPERATION AND MAINTENANCE STAGE"]
        for i, line in enumerate(lines):
            run = etree.SubElement(new_p, f"{{{W}}}r")
            rPr = etree.SubElement(run, f"{{{W}}}rPr")
            etree.SubElement(rPr, f"{{{W}}}b")
            sz2 = etree.SubElement(rPr, f"{{{W}}}sz")
            sz2.set(f"{{{W}}}val", "24")
            fonts2 = etree.SubElement(rPr, f"{{{W}}}rFonts")
            fonts2.set(f"{{{W}}}ascii", "Arial")
            t = etree.SubElement(run, f"{{{W}}}t")
            t.text = line
            if i < len(lines) - 1:
                br_run = etree.SubElement(new_p, f"{{{W}}}r")
                etree.SubElement(br_run, f"{{{W}}}br")

def insert_rsa_summary_table(source_doc, target_doc, marker_text: str):
    tables = source_doc.tables
    if not tables:
        return
    merged = [tables[-1]]
    for i in range(len(tables) - 2, -1, -1):
        if len(tables[i].columns) == len(merged[0].columns):
            merged.insert(0, tables[i])
        else: break
    base_el = deepcopy(merged[0]._element)
    for tbl in merged[1:]:
        for row in tbl.rows:
            base_el.append(deepcopy(row._element))
    _fix_run_spacing(base_el)
    _normalize_table_width(base_el, target_doc)
    _fix_header_row_text(base_el)
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None: return
    parent.insert(index, base_el)
    _remove_marker(parent, marker_el)


# ==========================================================
# FULL DOCUMENT INSERT (Annexures A/B/C)
# ==========================================================
def insert_full_document(source_doc, target_doc, marker_text: str, center_tables: bool = False):
    """
    Inserts a full document at the marker, preserving images and 
    handling page breaks cleanly without corrupting the master layout.
    """
    parent, index, marker_el = _find_marker(target_doc, marker_text)
    if parent is None:
        return

    for element in source_doc.element.body:
        # Skip main document section properties
        if element.tag == f"{{{W}}}sectPr":
            continue
            
        new_block = _transfer_image_rels(source_doc, target_doc, element)
        _strip_sect_pr(new_block)
        _preserve_xml_spaces(new_block)
            
        if center_tables:
            _center_tables_in_block(new_block)
        
        parent.insert(index, new_block)
        index += 1

    _remove_marker(parent, marker_el)
    log.info("Full document inserted at '%s'.", marker_text)


# ==========================================================
# ANNEXURE D — Extract from heading to end
# ==========================================================
def extract_till_end(source_doc, start_heading: str) -> list:
    content = []
    capture = False
    for tag, block in _iter_body(source_doc):
        if tag == "p" and start_heading in _block_text(block):
            capture = True
            continue
        if capture:
            content.append(deepcopy(block))
    return content

def insert_section_blocks(doc, marker_text: str, content_blocks: list, source_doc, image_scale: float = 1.15, table_scale: float = 1.1):
    parent, index, marker_el = _find_marker(doc, marker_text)
    if parent is None:
        return

    inserted_count = 0
    for block in content_blocks:
        tag = block.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            new_block = _transfer_image_rels(source_doc, doc, block)
            if tag == "tbl":
                _scale_table_in_block(new_block, scale=table_scale)
            elif tag == "p":
                _scale_images_in_block(new_block, scale=image_scale)
            
            _strip_sect_pr(new_block)
            _preserve_xml_spaces(new_block)

            # Anti-Blank-Page Logic
            if tag == "p":
                text_content = _block_text(new_block).strip()
                has_media = any(node.tag.endswith(('drawing', 'pict', 'object')) for node in new_block.iter())
                has_page_break = any(node.tag.endswith('br') and node.get(f"{{{W}}}type") == "page" for node in new_block.iter())
                
                if not text_content and not has_media and not has_page_break:
                    continue

            parent.insert(index, new_block)
            index += 1
            inserted_count += 1

    _remove_marker(parent, marker_el)


# ==========================================================
# ZONE COLOR
# ==========================================================
def apply_zone_color(doc, zone: str):
    color = ZONE_COLORS.get(zone)
    if not color:
        return
    for para in doc.paragraphs:
        for run in para.runs:
            if zone in run.text:
                run.font.color.rgb = color
                run.bold = True