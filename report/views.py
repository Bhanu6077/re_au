import logging
import os
import re
import uuid
import io
from datetime import datetime

from django.conf import settings
from django.http import FileResponse, HttpResponse
from django.shortcuts import render
from django.views import View

from docx import Document
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Cm

from .utils import (
    insert_road_table, copy_executive_summary_table,
    insert_section, insert_section_tables_only,
    copy_first_table_after_heading, copy_first_image_after_main_heading,
    copy_graph_after_table, insert_rsa_summary_table,
    insert_full_document, extract_till_end,
    insert_section_blocks, apply_zone_color,
    extract_executive_summary_values,
)

log = logging.getLogger(__name__)


def sanitize_filename(name: str) -> str:
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name)
    name = name.strip(". ")
    max_len = settings.MAX_FILENAME_LENGTH
    return (name[:max_len] if len(name) > max_len else name) or "Report"


def save_upload(django_file, prefix: str) -> str:
    ext = os.path.splitext(django_file.name)[1] or ".docx"
    dest = os.path.join(str(settings.UPLOAD_FOLDER), f"{prefix}_{uuid.uuid4().hex}{ext}")
    with open(dest, "wb") as f:
        for chunk in django_file.chunks():
            f.write(chunk)
    return dest


def cleanup(*paths):
    for p in paths:
        try:
            if p and os.path.exists(p):
                os.remove(p)
        except Exception as e:
            log.warning("Cleanup failed for %s: %s", p, e)


def parse_road_data(post):
    selected_types = post.getlist("road_type")
    road_data, road_summary = {}, []
    for road_type in selected_types:
        lhs = int(post.get(f"{road_type}_LHS", 0) or 0)
        rhs = int(post.get(f"{road_type}_RHS", 0) or 0)
        road_data[road_type] = {"LHS": lhs, "RHS": rhs}
        if lhs + rhs > 0:
            road_summary.append({"name": road_type, "total": lhs + rhs})
    return road_data, road_summary


class ReportView(View):

    def get(self, request):
        return render(request, "form.html")

    def post(self, request):
        temp_files = []
        try:
            tpl = DocxTemplate(str(settings.TEMPLATE_PATH))

            # DATE
            raw_date = request.POST.get("starting_survey_date", "")
            try:
                formatted_date = datetime.strptime(raw_date, "%Y-%m-%d").strftime("%d-%m-%Y")
            except ValueError:
                formatted_date = raw_date

            # MAP IMAGE
            image = None
            image_file = request.FILES.get("map_image")
            if image_file and image_file.name:
                image_path = save_upload(image_file, "map")
                temp_files.append(image_path)
                image = InlineImage(tpl, image_path, width=Cm(12))

            # ROAD DATA
            road_data, road_summary = parse_road_data(request.POST)

            # ANALYSED DOCUMENT — load early so we can extract template values
            analysed_path = None
            source_doc = None
            analysed_file = request.FILES.get("analysed_doc")
            if analysed_file and analysed_file.name:
                analysed_path = save_upload(analysed_file, "analysed")
                temp_files.append(analysed_path)
                source_doc = Document(analysed_path)

            # Extract executive-summary cell values for docxtpl placeholders
            exec_summary_values = (
                extract_executive_summary_values(source_doc)
                if source_doc else {}
            )

            # RENDER TEMPLATE
            temp_path = os.path.join(str(settings.UPLOAD_FOLDER), f"rendered_{uuid.uuid4().hex}.docx")
            temp_files.append(temp_path)

            context = {
                "project_name": request.POST.get("project_name", ""),
                "upc_code": request.POST.get("upc_code", ""),
                "state": request.POST.get("state", ""),
                "ro": request.POST.get("ro", ""),
                "piu": request.POST.get("piu", ""),
                "length": request.POST.get("length", ""),
                "flexibleorrigid": request.POST.get("flexibleorrigid", ""),
                "lanes": request.POST.get("lanes", ""),
                "om_dlp": request.POST.get("om_dlp", ""),
                "starting_survey_date": formatted_date,
                "zone": request.POST.get("zone", ""),
                "map_image": image,
                "road_summary": road_summary,
                # Executive summary table values (extracted from analysed_doc)
                **exec_summary_values,
            }

            tpl.render(context)
            tpl.save(temp_path)
            doc = Document(temp_path)

            # ANALYSED DOCUMENT — doc-manipulation pass (source_doc already loaded above)
            if source_doc is not None:
                anx_d_blocks = extract_till_end(source_doc, "Chainage Wise Gap Analysis")
                insert_section_blocks(doc, "### ANX_D ###", anx_d_blocks, source_doc)
                copy_executive_summary_table(source_doc, doc, "### INSERT_EXECUTIVE_SUMMARY_TABLE ###")
                # insert_section(doc, source_doc, "### INSERT_INVENTORY_SECTION_ONE ###", "3.1", "3.2")
                # insert_section(doc, source_doc, "### INSERT_INVENTORY_SECTION_TWO ###", "3.2", "3.3")

                insert_section_tables_only(doc, source_doc, "### INSERT_INVENTORY_SECTION_ONE ###", "3.1", "3.2")
                insert_section_tables_only(doc, source_doc, "### INSERT_INVENTORY_SECTION_TWO ###", "3.2", "3.3")

                insert_section(doc, source_doc, "### RESULT_GAP_STUDY_ONE ###", "4.1", "4.2")
                insert_section_tables_only(doc, source_doc, "### RESULT_GAP_STUDY_TWO ###", "4.2", "4.3")
                copy_first_table_after_heading(source_doc, doc, "### RESULT_GAP_STUDY_TWO_TABLE ###", "4.2")
                copy_first_image_after_main_heading(source_doc, doc, "### RESULT_GAP_STUDY_TWO_GRAPH_ONE ###", "Chainage Wise Gap Analysis")
                copy_graph_after_table(source_doc, doc, "### RESULT_GAP_STUDY_TWO_GRAPH_TWO ###", "4.2")

            # ANNEXURE A
            anx_a = request.FILES.get("anx_a_doc")
            if anx_a and anx_a.name:
                p = save_upload(anx_a, "anx_a"); temp_files.append(p)
                insert_full_document(Document(p), doc, "### ANX_A ###")

            # ANNEXURE B
            anx_b = request.FILES.get("anx_b_doc")
            if anx_b and anx_b.name:
                p = save_upload(anx_b, "anx_b"); temp_files.append(p)
                insert_full_document(Document(p), doc, "### ANX_B ###")

            # ANNEXURE C
            anx_c = request.FILES.get("anx_c_doc")
            if anx_c and anx_c.name:
                p = save_upload(anx_c, "anx_c"); temp_files.append(p)
                anx_c_doc = Document(p)
                # Yahaan center_tables=True pass kiya hai taaki sirf C ki tables theek hon
                insert_full_document(anx_c_doc, doc, "### ANX_C ###", center_tables=True)
                insert_rsa_summary_table(anx_c_doc, doc, "### RSA_SUMMARY ###")

            # ROAD TABLE & ZONE COLOR
            if road_data:
                insert_road_table(doc, "### INSERT_ROAD_TABLE ###", road_data)
            apply_zone_color(doc, request.POST.get("zone", ""))

            # SAVE OUTPUT
            safe_name = sanitize_filename(f"Final_Report_{request.POST.get('project_name', 'Report')}")
            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)
            log.info("Report generated in memory: %s.docx", safe_name)

            cleanup(*temp_files)

            return FileResponse(
                output_stream,
                as_attachment=True,
                filename=f"{safe_name}.docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        except Exception as e:
            log.exception("Report generation failed: %s", e)
            return HttpResponse(f"<h3>Error generating report</h3><pre>{e}</pre>", status=500)
        finally:
            cleanup(*temp_files)